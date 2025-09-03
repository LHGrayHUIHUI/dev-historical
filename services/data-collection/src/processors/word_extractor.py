"""
Word文档文本提取器

支持从Word文档中提取文本内容
"""

import logging
from typing import Dict, List, Any

from .base import TextExtractor

logger = logging.getLogger(__name__)


class WordExtractor(TextExtractor):
    """Word文档文本提取器
    
    支持从.docx和.doc格式的Word文档中提取文本
    """
    
    SUPPORTED_TYPES = {
        'application/msword',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'application/vnd.ms-word'
    }
    
    def __init__(self, config: Dict[str, Any] = None):
        """初始化Word提取器
        
        Args:
            config: 配置参数
        """
        super().__init__(config)
        self.extract_headers = config.get('extract_headers', True)
        self.extract_tables = config.get('extract_tables', True)
        self.preserve_formatting = config.get('preserve_formatting', False)
    
    def supports_file_type(self, file_type: str) -> bool:
        """检查是否支持指定文件类型"""
        return file_type in self.SUPPORTED_TYPES
    
    async def extract(self, file_path: str, **kwargs) -> List[Dict[str, Any]]:
        """从Word文档提取文本内容
        
        Args:
            file_path: Word文档路径
            **kwargs: 其他参数
            
        Returns:
            提取的文本内容列表
        """
        if not self.validate_file(file_path):
            raise ValueError(f"文件验证失败: {file_path}")
        
        try:
            # 检测文件类型并选择对应的提取方法
            if file_path.lower().endswith('.docx'):
                contents = await self._extract_from_docx(file_path)
            else:
                contents = await self._extract_from_doc(file_path)
            
            self.logger.info(f"Word文档提取完成: {file_path}, 段落数: {len(contents)}")
            return contents
            
        except Exception as e:
            self.logger.error(f"Word文档提取失败: {file_path}, 错误: {str(e)}")
            raise
    
    async def _extract_from_docx(self, file_path: str) -> List[Dict[str, Any]]:
        """从.docx文档提取文本
        
        Args:
            file_path: .docx文件路径
            
        Returns:
            提取的文本内容列表
        """
        contents = []
        
        try:
            from docx import Document
            
            # 打开文档
            doc = Document(file_path)
            
            # 提取段落内容
            paragraph_contents = await self._extract_paragraphs(doc)
            contents.extend(paragraph_contents)
            
            # 提取表格内容（可选）
            if self.extract_tables:
                table_contents = await self._extract_tables_from_docx(doc)
                contents.extend(table_contents)
            
            # 提取页眉页脚（可选）
            if self.extract_headers:
                header_footer_contents = await self._extract_headers_footers(doc)
                contents.extend(header_footer_contents)
            
            return contents
            
        except ImportError:
            self.logger.error("python-docx 未安装，无法处理.docx文件")
            raise ImportError("需要安装 python-docx 库")
        except Exception as e:
            self.logger.error(f"处理.docx文件失败: {str(e)}")
            raise
    
    async def _extract_from_doc(self, file_path: str) -> List[Dict[str, Any]]:
        """从.doc文档提取文本（需要转换）
        
        Args:
            file_path: .doc文件路径
            
        Returns:
            提取的文本内容列表
        """
        self.logger.warning("不支持直接处理.doc格式，建议转换为.docx格式")
        
        try:
            # 尝试使用python-docx2txt（如果可用）
            import docx2txt
            
            text = docx2txt.process(file_path)
            if text and text.strip():
                content_item = self.create_text_content(
                    content=text,
                    page_number=1,
                    title="Document Content",
                    extraction_method='docx2txt'
                )
                return [content_item]
            
            return []
            
        except ImportError:
            self.logger.error("docx2txt 未安装，无法处理.doc文件")
            raise ImportError("需要安装 docx2txt 库处理.doc文件")
    
    async def _extract_paragraphs(self, doc) -> List[Dict[str, Any]]:
        """提取文档段落
        
        Args:
            doc: Word文档对象
            
        Returns:
            段落内容列表
        """
        contents = []
        
        for para_idx, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if not text:
                continue
            
            # 检测段落样式
            style_name = paragraph.style.name if paragraph.style else "Normal"
            
            # 判断是否为标题
            is_heading = style_name.startswith('Heading') or style_name.startswith('Title')
            
            # 创建内容项
            content_item = self.create_text_content(
                content=text,
                page_number=para_idx + 1,
                title=text[:50] + "..." if len(text) > 50 and is_heading else None,
                paragraph_index=para_idx,
                style_name=style_name,
                is_heading=is_heading
            )
            
            contents.append(content_item)
        
        return contents
    
    async def _extract_tables_from_docx(self, doc) -> List[Dict[str, Any]]:
        """从Word文档提取表格
        
        Args:
            doc: Word文档对象
            
        Returns:
            表格内容列表
        """
        contents = []
        
        for table_idx, table in enumerate(doc.tables):
            try:
                # 提取表格数据
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        row_data.append(cell_text)
                    table_data.append(row_data)
                
                # 格式化表格为文本
                table_text = self._format_table_as_text(table_data)
                
                if table_text:
                    content_item = self.create_text_content(
                        content=table_text,
                        page_number=table_idx + 1,
                        title=f"Table {table_idx + 1}",
                        table_index=table_idx,
                        content_type='table'
                    )
                    contents.append(content_item)
            
            except Exception as e:
                self.logger.warning(f"提取表格 {table_idx} 失败: {str(e)}")
                continue
        
        return contents
    
    async def _extract_headers_footers(self, doc) -> List[Dict[str, Any]]:
        """提取页眉页脚
        
        Args:
            doc: Word文档对象
            
        Returns:
            页眉页脚内容列表
        """
        contents = []
        
        try:
            # 提取页眉
            for section_idx, section in enumerate(doc.sections):
                # 提取页眉
                if section.header:
                    for para in section.header.paragraphs:
                        text = para.text.strip()
                        if text:
                            content_item = self.create_text_content(
                                content=text,
                                page_number=section_idx + 1,
                                title=f"Header - Section {section_idx + 1}",
                                content_type='header',
                                section_index=section_idx
                            )
                            contents.append(content_item)
                
                # 提取页脚
                if section.footer:
                    for para in section.footer.paragraphs:
                        text = para.text.strip()
                        if text:
                            content_item = self.create_text_content(
                                content=text,
                                page_number=section_idx + 1,
                                title=f"Footer - Section {section_idx + 1}",
                                content_type='footer',
                                section_index=section_idx
                            )
                            contents.append(content_item)
        
        except Exception as e:
            self.logger.warning(f"提取页眉页脚失败: {str(e)}")
        
        return contents
    
    def _format_table_as_text(self, table_data: List[List[str]]) -> str:
        """将表格数据格式化为文本
        
        Args:
            table_data: 表格数据
            
        Returns:
            格式化后的文本
        """
        if not table_data:
            return ""
        
        # 计算每列的最大宽度
        max_widths = []
        for row in table_data:
            for col_idx, cell in enumerate(row):
                if col_idx >= len(max_widths):
                    max_widths.append(0)
                max_widths[col_idx] = max(max_widths[col_idx], len(cell))
        
        # 格式化表格
        formatted_rows = []
        for row_idx, row in enumerate(table_data):
            # 格式化每行
            formatted_cells = []
            for col_idx, cell in enumerate(row):
                width = max_widths[col_idx] if col_idx < len(max_widths) else 10
                formatted_cells.append(cell.ljust(width))
            
            formatted_row = " | ".join(formatted_cells)
            formatted_rows.append(formatted_row)
            
            # 在第一行后添加分隔线
            if row_idx == 0 and len(table_data) > 1:
                separator = " | ".join("-" * width for width in max_widths)
                formatted_rows.append(separator)
        
        return "\n".join(formatted_rows)
    
    def _detect_document_language(self, doc) -> str:
        """检测文档语言
        
        Args:
            doc: Word文档对象
            
        Returns:
            语言代码
        """
        # 收集一些文本样本进行语言检测
        sample_text = ""
        sample_count = 0
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip() and sample_count < 5:
                sample_text += paragraph.text + " "
                sample_count += 1
        
        if sample_text:
            return self.detect_language(sample_text)
        
        return None