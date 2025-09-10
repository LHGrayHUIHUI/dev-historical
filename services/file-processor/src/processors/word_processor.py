"""
Word文档处理器
专门处理Word文档(.doc/.docx)的文本提取和元数据获取
"""

import io
import logging
from typing import Dict, Any, Optional
from datetime import datetime
from pathlib import Path
import docx
from docx.opc.exceptions import PackageNotFoundError

logger = logging.getLogger(__name__)


class WordProcessor:
    """Word文档处理器
    
    提供Word文档(.docx)的文本提取、元数据获取等功能
    主要支持.docx格式，对.doc格式支持有限
    """
    
    def __init__(self):
        """初始化Word处理器"""
        self.supported_extensions = ['.docx', '.doc']
        logger.info("Word处理器初始化完成")
    
    async def process(
        self, 
        content: bytes, 
        extract_text: bool = True,
        extract_metadata: bool = True,
        filename: Optional[str] = None
    ) -> Dict[str, Any]:
        """处理Word文档
        
        Args:
            content: Word文件的二进制内容
            extract_text: 是否提取文本内容
            extract_metadata: 是否提取元数据
            filename: 文件名（可选）
            
        Returns:
            包含处理结果的字典
        """
        start_time = datetime.now()
        
        try:
            logger.info(f"开始处理Word文档，大小: {len(content)} bytes")
            
            result = {
                "filename": filename or "unknown.docx",
                "size": len(content),
                "format": "word",
                "text_content": "",
                "metadata": {},
                "paragraphs": 0,
                "processing_time": 0,
                "success": True,
                "warnings": []
            }
            
            # 使用BytesIO包装二进制内容
            doc_buffer = io.BytesIO(content)
            
            # 检查文件格式
            file_ext = Path(filename or "").suffix.lower()
            if file_ext == '.doc':
                result["warnings"].append("不完全支持.doc格式，建议转换为.docx格式")
                logger.warning("检测到.doc格式文件，可能处理效果不佳")
            
            # 提取文本内容
            if extract_text:
                try:
                    text_content, paragraph_count = await self._extract_text(doc_buffer)
                    result["text_content"] = text_content
                    result["paragraphs"] = paragraph_count
                    logger.info(f"成功提取文本，段落数: {paragraph_count}，文本长度: {len(text_content)}")
                    
                except Exception as e:
                    logger.error(f"文本提取失败: {e}")
                    result["text_content"] = ""
                    result["warnings"].append(f"文本提取失败: {str(e)}")
            
            # 提取元数据
            if extract_metadata:
                doc_buffer.seek(0)  # 重置缓冲区位置
                try:
                    metadata = await self._extract_metadata(doc_buffer)
                    result["metadata"] = metadata
                    logger.debug(f"成功提取元数据: {list(metadata.keys())}")
                    
                except Exception as e:
                    logger.warning(f"元数据提取失败: {e}")
                    result["warnings"].append(f"元数据提取失败: {str(e)}")
            
            # 计算处理时间
            processing_time = (datetime.now() - start_time).total_seconds()
            result["processing_time"] = round(processing_time, 3)
            
            logger.info(f"Word文档处理完成，处理时间: {processing_time:.3f}秒")
            
            return result
            
        except Exception as e:
            processing_time = (datetime.now() - start_time).total_seconds()
            logger.error(f"Word文档处理失败: {e}")
            
            return {
                "filename": filename or "unknown.docx",
                "size": len(content),
                "format": "word",
                "text_content": "",
                "metadata": {},
                "paragraphs": 0,
                "processing_time": round(processing_time, 3),
                "success": False,
                "error": str(e)
            }
    
    async def _extract_text(self, doc_buffer: io.BytesIO) -> tuple[str, int]:
        """提取Word文档的文本内容
        
        Args:
            doc_buffer: Word文件的BytesIO缓冲区
            
        Returns:
            tuple: (提取的文本内容, 段落数)
        """
        try:
            document = docx.Document(doc_buffer)
            
            text_parts = []
            paragraph_count = 0
            
            # 提取段落文本
            for paragraph in document.paragraphs:
                if paragraph.text.strip():  # 跳过空段落
                    text_parts.append(paragraph.text.strip())
                    paragraph_count += 1
            
            # 提取表格文本
            for table in document.tables:
                table_text = self._extract_table_text(table)
                if table_text:
                    text_parts.append(f"\n=== 表格内容 ===\n{table_text}")
            
            full_text = '\n\n'.join(text_parts)
            return full_text, paragraph_count
            
        except PackageNotFoundError:
            raise Exception("无效的Word文档格式或文件已损坏")
        except Exception as e:
            raise Exception(f"Word文档解析失败: {str(e)}")
    
    def _extract_table_text(self, table) -> str:
        """提取表格文本
        
        Args:
            table: python-docx的表格对象
            
        Returns:
            表格的文本内容
        """
        table_text = []
        
        try:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    row_text.append(cell_text if cell_text else " ")
                
                if any(cell.strip() for cell in row_text):  # 跳过完全空白的行
                    table_text.append(" | ".join(row_text))
            
            return '\n'.join(table_text)
            
        except Exception as e:
            logger.warning(f"表格文本提取失败: {e}")
            return ""
    
    async def _extract_metadata(self, doc_buffer: io.BytesIO) -> Dict[str, Any]:
        """提取Word文档元数据
        
        Args:
            doc_buffer: Word文件的BytesIO缓冲区
            
        Returns:
            包含元数据的字典
        """
        metadata = {}
        
        try:
            document = docx.Document(doc_buffer)
            
            # 文档核心属性
            core_props = document.core_properties
            
            # 提取标准元数据字段
            if core_props.title:
                metadata["title"] = core_props.title
            if core_props.author:
                metadata["author"] = core_props.author
            if core_props.subject:
                metadata["subject"] = core_props.subject
            if core_props.keywords:
                metadata["keywords"] = core_props.keywords
            if core_props.comments:
                metadata["comments"] = core_props.comments
            if core_props.category:
                metadata["category"] = core_props.category
            if core_props.language:
                metadata["language"] = core_props.language
            
            # 时间信息
            if core_props.created:
                metadata["created"] = core_props.created.isoformat()
            if core_props.modified:
                metadata["modified"] = core_props.modified.isoformat()
            if core_props.last_modified_by:
                metadata["last_modified_by"] = core_props.last_modified_by
            
            # 版本信息
            if core_props.revision:
                metadata["revision"] = str(core_props.revision)
            
            # 文档统计
            metadata["paragraphs"] = len(document.paragraphs)
            metadata["tables"] = len(document.tables)
            metadata["sections"] = len(document.sections)
            
            logger.debug(f"成功提取Word元数据字段: {list(metadata.keys())}")
            
        except Exception as e:
            logger.warning(f"Word元数据提取过程中发生错误: {e}")
            metadata["extraction_error"] = str(e)
        
        return metadata
    
    def is_supported(self, filename: str) -> bool:
        """检查文件是否支持
        
        Args:
            filename: 文件名
            
        Returns:
            是否支持该文件格式
        """
        file_ext = Path(filename).suffix.lower()
        return file_ext in self.supported_extensions